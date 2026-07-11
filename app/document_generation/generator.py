"""Document generation with Dashboard business-metrics synchronization."""

from __future__ import annotations

from datetime import date
import logging
from pathlib import Path
import shutil

from docx import Document
from openpyxl import Workbook

from app.company.profile import current_company_profile
from app.config.settings import get_settings
from app.config.user_settings import UserSettingsStore
from app.repositories.business_metrics import (
    BusinessMetricsRepository,
    BusinessStatus,
)


LOGGER = logging.getLogger(__name__)


class DocumentGenerator:
    def __init__(self) -> None:
        self.company = current_company_profile()
        prefs = UserSettingsStore().load()
        self.template_dir = (
            Path(prefs.template_dir)
            if prefs.template_dir
            else Path(__file__).resolve().parents[2]
            / "templates"
            / "company"
        )
        self.business_metrics = BusinessMetricsRepository()

    def output_dir(self, tender_id: int) -> Path:
        path = get_settings().data_dir / "outputs" / str(tender_id)
        path.mkdir(parents=True, exist_ok=True)
        return path

    def _replace(
        self,
        doc: Document,
        values: dict[str, str],
    ) -> None:
        def apply(paragraph) -> None:
            full = "".join(run.text for run in paragraph.runs)
            changed = full
            for key, value in values.items():
                changed = changed.replace(key, value)
            if changed == full:
                return
            if paragraph.runs:
                paragraph.runs[0].text = changed
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.text = changed

        for paragraph in doc.paragraphs:
            apply(paragraph)
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                apply(paragraph)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            apply(paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        apply(paragraph)

    def commercial_proposal(
        self,
        tender_id: int,
        report: dict,
        estimate: dict,
    ) -> Path:
        path = (
            self.output_dir(tender_id)
            / "Коммерческое_предложение.docx"
        )
        template = (
            self.template_dir
            / "01_Коммерческое_предложение.docx"
        )
        doc = Document(template) if template.exists() else Document()

        values = {
            "[ТЕКУЩАЯ ДАТА]": date.today().strftime("%d.%m.%Y"),
            "[АВТОМАТИЧЕСКИЙ НОМЕР]": f"КП-{tender_id:05d}",
            "[НАИМЕНОВАНИЕ ЗАКАЗЧИКА]": (
                report["tender"].get("customer") or "Не указан"
            ),
            "[НОМЕР И НАЗВАНИЕ ЗАКУПКИ]": (
                f"{report['tender'].get('number', '')} "
                f"{report['tender']['title']}"
            ),
            "[НАИМЕНОВАНИЕ И АДРЕС ОБЪЕКТА]": (
                report["tender"]["title"]
            ),
            "[СУММА]": f"{estimate['total']:,.2f}",
            "[СТАВКА]": str(estimate["vat_percent"]),
            "[СУММА НДС]": f"{estimate['vat']:,.2f}",
            "[СРОК]": "согласно документации закупки",
            "[УСЛОВИЯ]": (
                "согласно проекту договора после проверки рисков"
            ),
        }
        self._replace(doc, values)
        doc.add_heading("Расчёт стоимости", level=1)

        table = doc.add_table(rows=1, cols=5)
        headings = [
            "Наименование",
            "Количество",
            "Ед.",
            "Цена без НДС",
            "Сумма без НДС",
        ]
        for index, text in enumerate(headings):
            table.rows[0].cells[index].text = text

        for item in estimate["items"]:
            cells = table.add_row().cells
            cells[0].text = item["name"]
            cells[1].text = str(item["quantity"])
            cells[2].text = item["unit"]
            cells[3].text = (
                f"{item['price'] / item['quantity']:,.2f}"
                if item["quantity"]
                else "0"
            )
            cells[4].text = f"{item['price']:,.2f}"

        doc.add_paragraph(
            "Минимальная безопасная цена с НДС: "
            f"{estimate['minimum_safe_price_with_vat']:,.2f} руб."
        )
        doc.add_paragraph(
            f"Расчётная прибыль: {estimate['profit']:,.2f} руб.; "
            "рентабельность по выручке: "
            f"{estimate['margin_percent']:.2f}%."
        )
        doc.add_paragraph(
            f"Рекомендация системы: {report['recommendation']}"
        )
        doc.save(path)

        self._record_business_outputs(
            tender_id=tender_id,
            report=report,
            estimate=estimate,
            proposal_path=path,
        )
        return path

    def compliance_table(
        self,
        tender_id: int,
        report: dict,
    ) -> Path:
        path = (
            self.output_dir(tender_id)
            / "Таблица_соответствия.xlsx"
        )
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "Соответствие"
        worksheet.append(
            [
                "Требование заказчика",
                "Предлагаемое решение",
                "Соответствие",
                "Подтверждение",
                "Комментарий",
            ]
        )

        for equipment in report.get("equipment", []):
            worksheet.append(
                [
                    equipment["name"],
                    "Подлежит подбору, "
                    f"{equipment['quantity']} {equipment['unit']}",
                    "Требуется проверка",
                    "Паспорт/сертификат",
                    "Не предлагать несоответствующие аналоги",
                ]
            )

        if not report.get("equipment"):
            worksheet.append(
                [
                    "Перечень оборудования",
                    "В документации недостаточно информации",
                    "Требуется уточнение",
                    "",
                    "Направить запрос на разъяснение",
                ]
            )

        worksheet.freeze_panes = "A2"
        for width, column in zip(
            [35, 45, 22, 28, 42],
            "ABCDE",
            strict=True,
        ):
            worksheet.column_dimensions[column].width = width

        workbook.save(path)
        return path

    def clarification_request(
        self,
        tender_id: int,
        report: dict,
    ) -> Path:
        path = (
            self.output_dir(tender_id)
            / "Запрос_на_разъяснение.docx"
        )
        template = (
            self.template_dir
            / "03_Запрос_на_разъяснение.docx"
        )
        doc = Document(template) if template.exists() else Document()
        self._replace(
            doc,
            {
                "[ТЕКУЩАЯ ДАТА]": date.today().strftime(
                    "%d.%m.%Y"
                ),
                "[АВТОМАТИЧЕСКИЙ НОМЕР]": (
                    f"ЗР-{tender_id:05d}"
                ),
                "[НАИМЕНОВАНИЕ ЗАКАЗЧИКА]": (
                    report["tender"].get("customer")
                    or "Не указан"
                ),
                "[НОМЕР И НАЗВАНИЕ ЗАКУПКИ]": (
                    f"{report['tender'].get('number', '')} "
                    f"{report['tender']['title']}"
                ),
                "[НАИМЕНОВАНИЕ И АДРЕС ОБЪЕКТА]": (
                    report["tender"]["title"]
                ),
            },
        )

        risks = (
            report.get("license_requirements", [])
            + report.get("experience_requirements", [])
            + report.get("legal_risks", [])
            + report.get("competition_risks", [])
        )
        doc.add_heading(
            "Автоматически выявленные вопросы",
            level=1,
        )
        if not risks:
            doc.add_paragraph(
                "Автоматически значимые вопросы не выявлены. "
                "Перед направлением требуется ручная проверка "
                "полного комплекта документации."
            )
        for index, risk in enumerate(risks, 1):
            doc.add_paragraph(
                f"{index}. Просим разъяснить требование "
                f"«{risk['name']}». Найденный фрагмент: "
                f"{risk.get('quote', '')[:350]}"
            )

        doc.save(path)
        return path

    def package(
        self,
        tender_id: int,
        report: dict,
        estimate: dict,
    ) -> Path:
        output = self.output_dir(tender_id)
        self.commercial_proposal(tender_id, report, estimate)
        self.compliance_table(tender_id, report)
        self.clarification_request(tender_id, report)
        archive = shutil.make_archive(
            str(output / f"Пакет_заявки_{tender_id}"),
            "zip",
            output,
        )
        return Path(archive)

    def _record_business_outputs(
        self,
        *,
        tender_id: int,
        report: dict,
        estimate: dict,
        proposal_path: Path,
    ) -> None:
        try:
            tender_title = str(
                report.get("tender", {}).get(
                    "title",
                    "Тендер",
                )
            )
            self.business_metrics.record_estimate(
                tender_id,
                estimate,
                status=BusinessStatus.APPROVED,
                title=f"Смета — {tender_title}",
            )
            self.business_metrics.record_proposal(
                tender_id,
                file_path=proposal_path,
                total=estimate.get("total", 0),
                profit=estimate.get("profit", 0),
                status=BusinessStatus.READY,
                title=f"КП — {tender_title}",
            )
        except Exception:
            LOGGER.exception(
                "Не удалось синхронизировать КП и смету "
                "с Dashboard KPI"
            )
