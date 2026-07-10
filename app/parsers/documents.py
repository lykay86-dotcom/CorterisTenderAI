from __future__ import annotations
from pathlib import Path
from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook
from app.parsers.ocr import recognize_image, OCRUnavailable

SUPPORTED={".pdf",".docx",".xlsx",".xlsm",".txt",".csv",".jpg",".jpeg",".png",".tif",".tiff"}

def parse_document(path:Path)->tuple[str,int]:
    ext=path.suffix.lower()
    if ext==".pdf":
        reader=PdfReader(str(path)); chunks=[]
        for i,p in enumerate(reader.pages,1): chunks.append(f"\n[СТРАНИЦА {i}]\n{p.extract_text() or ''}")
        return "".join(chunks),len(reader.pages)
    if ext==".docx":
        doc=DocxDocument(str(path)); chunks=[p.text for p in doc.paragraphs]
        for table in doc.tables:
            chunks.extend(" | ".join(c.text for c in row.cells) for row in table.rows)
        return "\n".join(chunks),1
    if ext in {".xlsx",".xlsm"}:
        wb=load_workbook(path,read_only=True,data_only=True); chunks=[]
        for ws in wb.worksheets:
            chunks.append(f"[ЛИСТ {ws.title}]")
            for row in ws.iter_rows(values_only=True):
                vals=[str(v) for v in row if v is not None]
                if vals: chunks.append(" | ".join(vals))
        return "\n".join(chunks),len(wb.sheetnames)
    if ext in {".txt",".csv"}: return path.read_text(encoding="utf-8",errors="ignore"),1
    if ext in {".jpg",".jpeg",".png",".tif",".tiff"}:
        try:
            return recognize_image(path),1
        except OCRUnavailable as exc:
            return f"[OCR НЕДОСТУПЕН] {exc}",1
    return "",0

def classify_document(name:str,text:str)->str:
    s=(name+" "+text[:3000]).lower()
    rules=[("Проект договора",["проект договора","контракт"]),("Техническое задание",["техническое задание","описание объекта закупки"]),("Смета",["локальная смета","сметный расчет","ведомость объемов"]),("Требования к заявке",["состав заявки","инструкция участнику"]),("Спецификация",["спецификация","перечень оборудования"])]
    for kind,words in rules:
        if any(w in s for w in words): return kind
    return "Прочее"
