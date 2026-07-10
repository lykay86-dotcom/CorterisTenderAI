from __future__ import annotations

from dataclasses import dataclass, asdict
from pathlib import Path
import zipfile
from docx import Document
from openpyxl import load_workbook

@dataclass(slots=True)
class Check:
    name: str
    ok: bool
    severity: str
    details: str = ""


def inspect_docx(path: Path) -> list[Check]:
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    checks = [Check(f"{path.name}: документ не пустой", bool(text.strip() or doc.tables), "critical")]
    comments = path.with_suffix(path.suffix + ".comments")
    checks.append(Check(f"{path.name}: отсутствуют внешние комментарии", not comments.exists(), "warning"))
    return checks


def inspect_xlsx(path: Path) -> list[Check]:
    wb = load_workbook(path, read_only=False, data_only=False)
    hidden = [ws.title for ws in wb.worksheets if ws.sheet_state != "visible"]
    nonempty = any(ws.max_row > 1 or ws.max_column > 1 for ws in wb.worksheets)
    return [
        Check(f"{path.name}: книга не пустая", nonempty, "critical"),
        Check(f"{path.name}: отсутствуют скрытые листы", not hidden, "warning", ", ".join(hidden)),
    ]


def check_package(files: list[Path], max_size_mb: float = 50) -> dict:
    checks: list[Check] = []
    existing = [Path(x) for x in files if Path(x).exists()]
    checks.append(Check("Есть сформированные документы", bool(existing), "critical"))
    for path in existing:
        if path.suffix.lower() == ".docx": checks.extend(inspect_docx(path))
        if path.suffix.lower() in {".xlsx", ".xlsm"}: checks.extend(inspect_xlsx(path))
    size = sum(x.stat().st_size for x in existing) / 1024 / 1024
    checks.append(Check("Размер пакета не превышает лимит", size <= max_size_mb, "critical", f"{size:.2f} МБ"))
    critical_failed = any(not x.ok and x.severity == "critical" for x in checks)
    return {"ready": not critical_failed, "status": "Заявка готова к подаче" if not critical_failed else "Заявка не готова", "checks": [asdict(x) for x in checks]}
